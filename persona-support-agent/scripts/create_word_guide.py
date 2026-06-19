import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading_styled(doc, text, level, color_rgb):
    """Adds a styled heading with custom font size and color."""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.color.rgb = color_rgb
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    return h

def main():
    doc = Document()
    
    # Configure margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Theme Colors
    indigo = RGBColor(79, 70, 229)    # #4F46E5 - Titles
    slate = RGBColor(30, 41, 59)      # #1E293B - Headings
    charcoal = RGBColor(51, 65, 85)   # #334155 - Sub-sections
    body_color = RGBColor(71, 85, 105)# #475569 - Body text

    # Set normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = body_color

    # Cover Header / Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("Adsparkx AI Agent Console\n")
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = indigo
    
    run_subtitle = title_p.add_run("Persona-Aware Support Console - Screen Recording & Submission Guide\n")
    run_subtitle.font.size = Pt(14)
    run_subtitle.font.bold = True
    run_subtitle.font.color.rgb = slate
    
    run_meta = title_p.add_run("A complete step-by-step presentation script covering all assignment requirements.")
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(148, 163, 184)
    title_p.paragraph_format.space_after = Pt(30)

    # Section 1: Recording Overview
    add_heading_styled(doc, "1. Video Presentation Overview", 1, slate)
    p = doc.add_paragraph(
        "This guide provides a professional narration script and visual layout instructions for recording your "
        "assignment video submission. Your video should be between 3 to 8 minutes long and demonstrate a thorough "
        "understanding of both the user-facing app capabilities and the underlying code architecture. Ensure you have "
        "the Streamlit app running locally (on port 8501) and have a clean console terminal ready before recording."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)

    # Section 2: Requirement Check
    add_heading_styled(doc, "2. Mapping Presentation Sections to Evaluation Criteria", 1, slate)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Recording Target'
    hdr_cells[1].text = 'Technical Implementation File'
    hdr_cells[2].text = 'Demonstration Steps'
    
    for cell in hdr_cells:
        set_cell_background(cell, '4F46E5')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True
        
    items = [
        ("Project Structure Overview", "Root / Modules layout", "Show project folder inside VS Code or explorer, highlighting modular components."),
        ("Knowledge Base Ingestion", "scripts/generate_data.py & src/rag_pipeline.py", "Show document loading, splitting, persistent ChromaDB index, and embedding fallback."),
        ("Persona Detection", "src/classifier.py", "Demonstrate classification of Technical, Frustrated, and Executive personas."),
        ("Semantic Retrieval", "src/rag_pipeline.py", "Point out retrieved chunks, match scores, and progress meter bars on the console."),
        ("Adaptive Persona Generation", "src/generator.py", "Demonstrate how responses adapt tone, details, and format matching the user."),
        ("5 Different User Queries", "app.py User Interface", "Input 5 distinct queries showcasing RAG retrieval, fallback rules, and personas."),
        ("Escalation Demonstration", "src/escalator.py", "Trigger human handoff by entering a billing query or triggering low similarity scores."),
        ("Human Handoff JSON Report", "src/escalator.py", "Show the generated handoff JSON summary, listing metadata, sentiments, and history."),
        ("Technical Design Decision", "Architecture explanation", "Detail the offline fallback capability and hybrid API system design decision.")
    ]
    
    for target, source, steps in items:
        row_cells = table.add_row().cells
        row_cells[0].text = target
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = source
        row_cells[2].text = steps
        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Detailed Script
    add_heading_styled(doc, "3. Script & Walkthrough", 1, slate)

    # 3.1 Structure
    add_heading_styled(doc, "Step 1: Project Structure Overview (Duration: ~30s)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"Hello and welcome! In this presentation, I am going to demonstrate the Adsparkx AI Adaptive Support Agent Console, "
        "built to dynamically classify customer personas, perform context-grounded retrieval-augmented generation (RAG) using ChromaDB, "
        "and escalate to human agents under specific safety rules. "
        "Let's start by looking at our codebase structure. "
        "As you can see on screen, the project is structured modularly: \n"
        "- The 'app.py' file runs our custom Streamlit front-end console. \n"
        "- The 'cli.py' file provides an alternative interactive terminal CLI chatbot. \n"
        "- Inside 'src/', we have separated concerns: 'config.py' handles custom thresholds; 'classifier.py' manages persona detection; "
        "'rag_pipeline.py' chunk-splits our corpus and queries the vector store; 'generator.py' compiles adaptive persona prompts; "
        "and 'escalator.py' checks handoff conditions.\n"
        "- In 'scripts/generate_data.py', we generate our text documents and a PDF support guide.\""
    )

    # 3.2 Ingestion
    add_heading_styled(doc, "Step 2: Knowledge Ingestion Process (Duration: ~45s)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"To populate our RAG knowledge base, we run 'scripts/generate_data.py', which generates 13 MD/TXT articles and "
        "programmatically compiles a PDF support guide. "
        "If we open 'src/rag_pipeline.py', we can see that document ingestion uses LangChain's text splitters to split documents "
        "into 400-character chunks with an overlap. "
        "These chunks are processed using embeddings (Gemini's 'text-embedding-004' by default, falling back to a local "
        "'sentence-transformers/all-MiniLM-L6-v2' model if offline) and stored into a persistent ChromaDB database. "
        "Our persistent ChromaDB holds 68 unique document chunks, ready for semantic search.\""
    )

    # 3.3 Persona Detection
    add_heading_styled(doc, "Step 3: Persona Detection & Classification (Duration: ~45s)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"One of the core requirements is persona-aware classification. The system dynamically groups customers into: \n"
        "1. Technical Expert (expects precise parameters, diagnostic data, and configuration logs) \n"
        "2. Frustrated User (needs high empathy, reassurance, and polite tone) \n"
        "3. Business Executive (expects operational summaries, cost/business impact analysis, and resolution ETAs) \n"
        "If we open 'src/classifier.py', you can see that the detection is context-aware—we pass the rolling conversation history "
        "along with the query to Gemini Structured Outputs. If the Gemini API is offline, a robust rule-based classifier resolves "
        "the persona using custom regex word boundaries, ensuring the agent console is fully testable offline.\""
    )

    # 3.4 Retrieval
    add_heading_styled(doc, "Step 4: Semantic Retrieval Process (Duration: ~45s)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"When a user types a query, 'rag_pipeline.py' executes a semantic query against ChromaDB using cosine distance. "
        "We restrict retrieval to a similarity threshold configured in 'src/config.py'. "
        "On the right-hand Diagnostics panel of our Streamlit dashboard, you can see 'Retrieved Knowledge Base Sources' "
        "rendering a progress bar representing the similarity match level for each chunk. "
        "This gives developers full transparency into RAG scoring and helps verify grounding quality in real time.\""
    )

    # 3.5 Responses
    add_heading_styled(doc, "Step 5: Responses for All Three Personas (Duration: ~1m)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"Once a persona is classified and sources are retrieved, 'src/generator.py' compiles adaptive response prompts. "
        "If a Technical Expert is identified, the model responds with code blocks, troubleshooting logs, and header configs. "
        "For a Frustrated User, the agent uses empathetic openings, apologizes for delays, and prioritizes reassurance. "
        "For a Business Executive, it ignores raw codes and instead lists financial ETAs, business impacts, and clear summaries. "
        "Let's see this in action on our dashboard. I will type three queries to show tone adaptation: \n"
        "- For a technical user query: 'Explain the webhook payload signature validation headers.' The response includes code snippets and parameter breakdowns. \n"
        "- For a frustrated user query: 'This is completely broken! I need help now!' The response expresses deep empathy and reassurance. \n"
        "- For a business executive query: 'What is the business impact and resolution ETA?' The response is presented in clean bullet points focused on executive details.\""
    )

    # 3.6 Five queries
    add_heading_styled(doc, "Step 6: Demonstration of 5 User Queries (Duration: ~1m)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"To demonstrate the console, I will now type 5 different user queries. "
        "Notice how the diagnostics panel on the right updates classification confidence, reasoning, and sources for each turn: \n\n"
        "Query 1 (Technical): 'Explain the bearer token authorization headers' \n"
        "-> Show the Cyan badge, reasoning, and code blocks in response. \n\n"
        "Query 2 (Frustrated): 'Nothing is loading on your interface, it has been an hour!' \n"
        "-> Show the Rose badge, high empathy agent response, and matching sources. \n\n"
        "Query 3 (Executive): 'What are the subscription models and pricing tiers?' \n"
        "-> Show the Amber badge, structured pricing summaries, and no raw codes. \n\n"
        "Query 4 (General): 'How do I clear browser cookies?' \n"
        "-> Show standard grounding from our cookie guide article. \n\n"
        "Query 5 (Escalation Trigger): 'I have a billing issue with my invoice' \n"
        "-> Let's type this billing query, which leads to our next requirement: Escalation.\""
    )

    # 3.7 Escalation & JSON Handoff
    add_heading_styled(doc, "Step 7 & 8: Escalation & Human Handoff Summary (Duration: ~1m)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"If the system detects sensitive keywords (like 'sue', 'legal', 'billing', 'invoice'), a semantic search confidence score "
        "below the threshold (such as 0.40), an explicit user request for a human, or 3 consecutive turns of customer frustration, "
        "'src/escalator.py' immediately triggers human handoff. "
        "Let's watch this happen: as I type 'I have a billing issue', the escalation triggers. "
        "The chat interface immediately locks inputs and shows a red warning alert: 'This chat session has been escalated'. "
        "On the right diagnostics panel under 'Escalation Summary', a red glowing alert flashes. "
        "We also generate a structured Handoff JSON Report. If I expand it, you can see all required metadata: "
        "the customer persona, session history, sentiment, and the exact trigger reason. "
        "This JSON is saved in our logs directory to hand over to a human specialist.\""
    )

    # 3.8 Technical Design
    add_heading_styled(doc, "Step 9: Explanation of Technical Design Decision (Duration: ~45s)", 2, charcoal)
    doc.add_paragraph(
        "NARRATION:\n"
        "\"I will wrap up by explaining a key technical design decision: our Offline-First Fallback Architecture. "
        "To ensure the system works reliably under API rate limits or network issues, we designed a dual-mode pipeline. "
        "Our embeddings and classification subsystems automatically detect the presence of a valid Gemini API key. "
        "If online, we utilize Gemini's high-speed structured outputs and 'text-embedding-004' models. "
        "If offline, the system falls back seamlessly to a local 'sentence-transformers' model for embeddings and a regex-based keyword "
        "classifier for persona detection. "
        "This fallback design guarantees 100% testability and stability in any local development environment while still allowing "
        "us to leverage state-of-the-art LLM capabilities when connected. "
        "Thank you for watching this presentation!\""
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Section 4: Tips for Recording
    add_heading_styled(doc, "4. Tips for a High-Quality Screen Recording", 1, slate)
    
    tips = [
        "**Prepare the empty state:** Start the video with an empty conversation so you can showcase the modern welcome board and card suggestions.",
        "**Hover effects:** Move your cursor over the diagnostics cards to show off the glassmorphic translateY lift hover transitions.",
        "**Speak clearly:** Follow the narration scripts in this document closely, but feel free to speak naturally.",
        "**Keep code snippets open:** Have VS Code open to show `src/classifier.py` and `src/rag_pipeline.py` when discussing steps 2 and 3.",
        "**Check your API status:** In the Streamlit sidebar, verify your API connection status indicator so you know whether you are demoing the Gemini LLM or local keyword fallbacks."
    ]
    
    for tip in tips:
        p_tip = doc.add_paragraph(style='List Bullet')
        p_tip.paragraph_format.space_after = Pt(4)
        run_bold = p_tip.add_run(tip.split(":")[0] + ":")
        run_bold.bold = True
        p_tip.add_run(tip.split(":")[1])

    # Save document
    output_path = "c:\\Users\\User\\Desktop\\Assignment\\persona-support-agent\\screen_recording_guide.docx"
    doc.save(output_path)
    print(f"Document created successfully at: {output_path}")

if __name__ == "__main__":
    main()
